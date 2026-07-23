"""Plain-text extraction from .docx files — the base resume, and later
tailored resumes registered via `save-resume` (M7), both need their text
content for fit scoring."""
from __future__ import annotations

from pathlib import Path

import docx


def extract_docx_text(path: str | Path) -> str:
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)
